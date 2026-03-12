from docx import Document
from docx.shared import Pt

# Create a new Document
doc = Document()

# Function to add equations
def add_equation(equation):
    p = doc.add_paragraph()
    run = p.add_run(equation)
    run.font.size = Pt(12)
    run.font.name = 'Cambria Math'

# Add title
doc.add_heading('线性回归计算过程', level=1)

# Add problem description
doc.add_heading('题目描述', level=2)
problem_text = (
    "已知某公司广告费用（单位：万元）与销售额（单位：万元）的数据如下表所示：\n\n"
)
doc.add_paragraph(problem_text)

# Create the table with vertical columns
data = [
    ('广告费用 x（万元）', '销售额 y（万元）'),
    (2.0, 5.0),
    (2.5, 5.5),
    (3.0, 6.0),
    (3.5, 7.0),
    (4.0, 7.5),
    (4.5, 8.0),
    (5.0, 8.5),
    (5.5, 9.0),
    (6.0, 9.5),
    (6.5, 10.0)
]

table = doc.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '广告费用 x（万元）'
hdr_cells[1].text = '销售额 y（万元）'

for x, y in data[1:]:
    row_cells = table.add_row().cells
    row_cells[0].text = str(x)
    row_cells[1].text = str(y)

# Step 1: 计算平均值
doc.add_heading('Step 1: 计算平均值', level=2)
add_equation(r'广告费用平均值: \(\bar{x} = \frac{1}{n} \sum_{i=1}^n x_i = 4.25\)')
add_equation(r'销售额平均值: \(\bar{y} = \frac{1}{n} \sum_{i=1}^n y_i = 7.6\)')

# Step 2: 计算相关系数
doc.add_heading('Step 2: 计算相关系数', level=2)
add_equation(r'广告费用的方差: \(L_{xx} = \sum_{i=1}^n (x_i - \bar{x})^2 = 20.63\)')
add_equation(r'销售额的方差: \(L_{yy} = \sum_{i=1}^n (y_i - \bar{y})^2 = 26.40\)')
add_equation(r'协方差: \(L_{xy} = \sum_{i=1}^n (x_i - \bar{x})(y_i - \bar{y}) = 23.25\)')

# Step 3: 计算估计参数
doc.add_heading('Step 3: 计算估计参数', level=2)
add_equation(r'\hat{\beta}_1 = \frac{\sum (x_i - \bar{x})(y_i - \bar{y})}{\sum (x_i - \bar{x})^2} = \frac{23.25}{20.63} \approx 1.13')
add_equation(r'\hat{\beta}_0 = \bar{y} - \hat{\beta}_1 \bar{x} = 7.60 - 1.13 \times 4.25 \approx 2.81')

# Step 4: 直线方程
doc.add_heading('Step 4: 直线方程', level=2)
add_equation(r'L: \hat{y} = 1.13 x + 2.81')

# Step 6: 误差分解
doc.add_heading('Step 6: 误差分解', level=2)
add_equation(r'残差平方和: \(SS_e = L_{yy} - \hat{\beta}_1 L_{xy} = 26.40 - 1.13 \times 23.25 \approx 0.19\)')
add_equation(r'估计标准差: \(\hat{\sigma} = \sqrt{\frac{SS_e}{n - k - 1}} = \sqrt{\frac{0.19}{10 - 2}} \approx 0.15\)')
add_equation(r'相关系数: \(r = \frac{\sum (x_i - \bar{x})(y_i - \bar{y})}{\sqrt{\sum (x_i - \bar{x})^2 \sum (y_i - \bar{y})^2}} = \frac{23.25}{\sqrt{20.63 \times 26.40}} \approx 1.00\)')
add_equation(r'决定系数: \(R^2 = \frac{L^2_{xy}}{L_{xx} L_{yy}} = \frac{23.25^2}{20.63 \times 26.40} \approx 0.99\)')

# Save the document
file_path = "/Doc1.docx"
doc.save(file_path)

file_path